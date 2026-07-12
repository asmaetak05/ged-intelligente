import zipfile
import os
from docx import Document

os.makedirs("tests/fixtures", exist_ok=True)

doc = Document()
doc.add_paragraph("Document DOCX de test.")
doc.add_paragraph("Objet: Travaux de construction d'une route.")
doc.add_paragraph("Maitre d'ouvrage: Royaume du Maroc.")
doc.add_paragraph("Estimation: 50 000 MAD.")
doc.save("tests/fixtures/test_doc.docx")

with zipfile.ZipFile("tests/fixtures/sample_ao.zip", "w") as z:
    z.write("tests/fixtures/test_doc.docx", arcname="test_doc.docx")
    z.writestr("test_scan.pdf", "%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >> >>\nendobj\n4 0 obj\n<< /Length 53 >>\nstream\nBT\n/F1 24 Tf\n100 700 Td\n(Test PDF) Tj\nET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000288 00000 n \ntrailer\n<< /Size 5 /Root 1 0 R >>\nstartxref\n390\n%%EOF")
